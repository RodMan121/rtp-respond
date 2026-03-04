"""
word_export.py — Export de la réponse RFP en Word
==================================================
Produit un document .docx structuré :
- Page de garde
- Synthèse exécutive
- Réponse par exigence (groupée par type)
- Annexe : exigences implicites et prédictives
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from loguru import logger
from pathlib import Path

_MOSCOW_LABELS = {
    "must": "OBLIGATOIRE",
    "should": "RECOMMANDÉ",
    "could": "OPTIONNEL",
    "wont": "HORS PÉRIMÈTRE",
}

_TYPE_LABELS = {
    "functional":   "Exigences Fonctionnelles",
    "performance":  "Performance",
    "security":     "Sécurité",
    "operational":  "Exploitation & Maintenance",
    "data":         "Données",
    "compliance":   "Conformité",
    "usability":    "Ergonomie & Accessibilité",
    "implicit":     "Besoins Implicites",
    "predictive":   "Besoins Prédictifs",
}

def _set_heading_style(paragraph, level=1):
    paragraph.style = f"Heading {level}"

def _add_colored_run(paragraph, text: str, color: tuple, bold=False, size=10):
    run = paragraph.add_run(text)
    run.font.color.rgb = RGBColor(*color)
    run.font.bold = bold
    run.font.size = Pt(size)
    return run

def _add_page_break(doc):
    doc.add_paragraph().add_run().add_break(
        __import__("docx.enum.text", fromlist=["WD_BREAK"]).WD_BREAK.PAGE
    )

def _cover_page(doc: Document, rfp_name: str, total: int, must_count: int):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"\n\n\nRÉPONSE À L'APPEL D'OFFRES\n")
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = RGBColor(47, 79, 127)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(rfp_name)
    r2.font.size = Pt(16)
    r2.font.color.rgb = RGBColor(80, 80, 80)

    doc.add_paragraph("\n")
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(
        f"Matrice de réponse — {total} exigences analysées\n"
        f"dont {must_count} exigences obligatoires (Must)"
    )
    r3.font.size = Pt(11)
    r3.font.color.rgb = RGBColor(100, 100, 100)
    _add_page_break(doc)

def _executive_summary(doc: Document, matrix: list[dict]):
    _set_heading_style(doc.add_paragraph("Synthèse Exécutive"), level=1)

    # Comptages
    by_moscow = {m: sum(1 for x in matrix if x.get("moscow") == m)
                 for m in ["must", "should", "could", "wont"]}
    by_source = {s: sum(1 for x in matrix if x.get("source") == s)
                 for s in ["explicit", "implicit", "predictive", "signal"]}

    p = doc.add_paragraph()
    p.add_run(
        f"Notre analyse du document a permis d'identifier {len(matrix)} exigences, "
        f"dont {by_moscow['must']} exigences obligatoires, "
        f"{by_moscow['should']} exigences recommandées "
        f"et {by_moscow.get('could', 0)} exigences optionnelles.\n\n"
    ).font.size = Pt(10)

    p2 = doc.add_paragraph()
    p2.add_run(
        f"Au-delà des {by_source.get('explicit', 0)} exigences explicitement formulées, "
        f"notre analyse approfondie a révélé {by_source.get('implicit', 0) + by_source.get('predictive', 0)} "
        f"besoins implicites et prédictifs, ainsi que "
        f"{by_source.get('signal', 0)} signaux faibles révélateurs des enjeux prioritaires du projet.\n\n"
        f"Notre réponse couvre l'intégralité de ces besoins."
    ).font.size = Pt(10)
    _add_page_break(doc)

def _requirement_section(doc: Document, items: list[dict], type_label: str):
    if not items:
        return
    _set_heading_style(doc.add_paragraph(type_label), level=2)

    for item in items:
        moscow = item.get("moscow", "should")
        source = item.get("source", "explicit")

        # En-tête de l'exigence
        p_id = doc.add_paragraph()
        _add_colored_run(p_id, f"[{item['id']}]  ", (47, 79, 127), bold=True, size=10)
        _add_colored_run(p_id, _MOSCOW_LABELS.get(moscow, moscow), (80, 80, 80), size=9)
        if source in ("signal", "implicit", "predictive"):
            _add_colored_run(p_id, f"  ◆ {source.upper()}", (150, 80, 0), size=8)

        # Texte de l'exigence
        p_text = doc.add_paragraph()
        p_text.paragraph_format.left_indent = Cm(0.5)
        run_text = p_text.add_run(item.get("text", ""))
        run_text.font.size = Pt(10)
        run_text.font.bold = True

        # Citation source (si disponible)
        if item.get("source_quote"):
            p_quote = doc.add_paragraph()
            p_quote.paragraph_format.left_indent = Cm(0.5)
            r = p_quote.add_run(f"« {item['source_quote']} »")
            r.font.size = Pt(9)
            r.font.italic = True
            r.font.color.rgb = RGBColor(100, 100, 100)

        # Réponse proposée
        if item.get("proposed_response"):
            p_resp = doc.add_paragraph()
            p_resp.paragraph_format.left_indent = Cm(0.5)
            _add_colored_run(p_resp, "→ ", (47, 79, 127), bold=True, size=10)
            r_resp = p_resp.add_run(item["proposed_response"])
            r_resp.font.size = Pt(10)
            r_resp.font.color.rgb = RGBColor(30, 30, 30)

        # Justification (si implicite/prédictif)
        if item.get("rationale") and source in ("signal", "implicit", "predictive"):
            p_rat = doc.add_paragraph()
            p_rat.paragraph_format.left_indent = Cm(0.5)
            r_rat = p_rat.add_run(f"Justification : {item['rationale']}")
            r_rat.font.size = Pt(8)
            r_rat.font.italic = True
            r_rat.font.color.rgb = RGBColor(120, 80, 40)

        doc.add_paragraph()  # espacement

def export(matrix: list[dict], output_path: str, rfp_name: str = "RFP"):
    doc = Document()

    # Marges
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    must_count = sum(1 for x in matrix if x.get("moscow") == "must")
    _cover_page(doc, rfp_name, len(matrix), must_count)
    _executive_summary(doc, matrix)

    # Grouper par type, exigences explicites d'abord
    explicit = [x for x in matrix if x.get("source") in ("explicit", "signal", "functional", "constraints")]
    augmented = [x for x in matrix if x.get("source") in ("implicit", "predictive", "stakeholders")]

    # Section 1 : réponse aux exigences explicites
    _set_heading_style(doc.add_paragraph("1. Réponse aux Exigences"), level=1)
    types_explicit = sorted(set(x.get("type") for x in explicit))
    for t in types_explicit:
        items_t = [x for x in explicit if x.get("type") == t]
        _requirement_section(doc, items_t, _TYPE_LABELS.get(t, t))

    _add_page_break(doc)

    # Section 2 : valeur ajoutée (besoins implicites et prédictifs)
    _set_heading_style(doc.add_paragraph("2. Notre Valeur Ajoutée"), level=1)
    p_intro = doc.add_paragraph()
    p_intro.add_run(
        "Au-delà des exigences formellement exprimées, notre analyse a identifié "
        "les besoins suivants que nous prenons en charge proactivement."
    ).font.size = Pt(10)
    doc.add_paragraph()

    types_aug = sorted(set(x.get("type") for x in augmented))
    for t in types_aug:
        items_t = [x for x in augmented if x.get("type") == t]
        _requirement_section(doc, items_t, _TYPE_LABELS.get(t, t))

    doc.save(output_path)
    logger.success(f"📝 Export Word : {output_path}")
